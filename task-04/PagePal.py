from typing import Final;
import requests,csv 
from telegram import Update,InlineKeyboardMarkup,InlineKeyboardButton 
from telegram.ext import Application,CommandHandler,MessageHandler,filters,ContextTypes,CallbackQueryHandler 
import docx 

Token:Final='7407195288:AAG1LMZX3InXF--Q2rMdsCAAo9Gb8l_bFo0'
Username:Final='@PagePalAssistBot'
BooksAPI:Final='AIzaSyDZZwcAEPI0AlXtQK7eJ3jMWNzdi5muUTI'

async def start_command(update:Update,context:ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text('Hello Welcome to PagePal use /help command for more option.')

async def help_command(update:Update,context:ContextTypes.DEFAULT_TYPE):
    text=(
        "*Welcome to PagePal!*\n\n"
        "Here's what I can do for you:\n\n"
        "/book - Tell me a genre, and I'll fetch a list of books for you.\n\n"
        "/preview - Share a book name, and I'll find a preview link for you.\n\n"
        "/list - Type in a specific book name and manage your reading list.\n\n"
        "/reading_list - Manage your reading list:\n"
        "    - Add a book to reading list\n"
        "    - Delete a book from reading list\n"
        "    - View your entire reading list in a docx file\n"
    )
    await update.message.reply_text(text)

async def book_command(update:Update,context:ContextTypes.DEFAULT_TYPE):
    context.user_data['genre']=True
    await update.message.reply_text('Please enter a genre')

async def preview_command(update:Update,context:ContextTypes.DEFAULT_TYPE):
    context.user_data['book']=True
    await update.message.reply_text('Please enter the name of the book')

async def list_command(update:Update,context:ContextTypes.DEFAULT_TYPE):
    context.user_data['list']=True
    await update.message.reply_text('Please enter the name of the book')


async def reading_list_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    book=context.user_data.get('list_book')
    if not book:
        await update.message.reply_text('Please use the /list command first to select a book.')
        return
    buttons=[
        [InlineKeyboardButton("Add book to the Reading List",callback_data='add')],
        [InlineKeyboardButton("Delete book from the Reading List",callback_data='delete')],
        [InlineKeyboardButton("View Reading List",callback_data='view')]
    ]
    reply=InlineKeyboardMarkup(buttons)
    await update.message.reply_text('Manage your reading list:',reply_markup=reply)

reading_list=[]

async def reading_list_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query=update.callback_query
    book=context.user_data.get('list_book')
    await query.answer()
    if not book:
        await query.edit_message_text('No book selected. Please use the /list command first to select a book.')
        return
    if query.data=="add":
        if book not in reading_list:
            reading_list.append(book)
            await query.edit_message_text(text=f'"{book}" has been added to your reading list.')
        else:
            await query.edit_message_text(text=f'"{book}" is already in your reading list.')
    elif query.data=="delete":
        if book in reading_list:
            reading_list.remove(book)
            await query.edit_message_text(text=f'"{book}" has been removed from your reading list.')
        else:
            await query.edit_message_text(text=f'"{book}" is not in your reading list.')
    elif query.data=="view":
        if reading_list:
            document=docx.Document()
            document.add_heading('My Reading List',0)
            for book in reading_list:
                response=requests.get(f'https://www.googleapis.com/books/v1/volumes?q={book}&key={BooksAPI}&maxResults=1')
                data=response.json()
                if 'items' in data:
                    title=data['items'][0]['volumeInfo'].get('title','Unknown Title')
                    authors=', '.join(data['items'][0]['volumeInfo'].get('authors',['Unknown Author']))
                    publisher=data['items'][0]['volumeInfo'].get('publisher','N/A') 
                    date=data['items'][0]['volumeInfo'].get('publishedDate','N/A')  
                    preview=data['items'][0]['volumeInfo'].get('previewLink','N/A')
                    description=data['items'][0]['volumeInfo'].get('description','No description available.')
                    document.add_heading(title,level=1)
                    document.add_paragraph(f'Authors: {authors}')
                    document.add_paragraph(f'Published Date: {date}')
                    document.add_paragraph(f'Publisher: {publisher}')
                    document.add_paragraph(f'Preview: {preview}')
                    document.add_paragraph(f'Description: {description}')
                else:
                    document.add_heading(book,level=1)
                    document.add_paragraph('Preview not available.')
            filename='ReadingList.docx'
            document.save(filename)
            await query.message.reply_document(document=open(filename,'rb'))
        else:
            await query.edit_message_text('Your reading list is empty.')

async def book_handler(update:Update,context:ContextTypes.DEFAULT_TYPE):
    genre=update.message.text.lower()
    response=requests.get(f'https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={BooksAPI}')
    data=response.json()
    if 'items' in data:
        books=[]
        for book in data['items']:
            title=book['volumeInfo'].get('title','Unknown Title')
            authors=', '.join(book['volumeInfo'].get('authors',['Unknown Author']))
            publisher=book['volumeInfo'].get('publisher','N/A') 
            date=book['volumeInfo'].get('publishedDate','N/A')  
            preview=book['volumeInfo'].get('previewLink','N/A')
            description=book['volumeInfo'].get('description','No description available.')
            books.append([title,authors,publisher,date,preview,description])
        filename='BooksList.csv'
        with open(filename,'w') as file:
                writer=csv.writer(file)
                writer.writerow(['TITLE','AUTHORS','PUBLISHER','PUBLISHED DATE','PREVIEW LINK','DESCRIPTION'])
                writer.writerows(books)
        await update.message.reply_document(document=open(filename,'rb'))
    else:
        await update.message.reply_text('No books found.')

async def preview_handler(update:Update,context:ContextTypes.DEFAULT_TYPE):
    book=update.message.text.lower()
    response=requests.get(f'https://www.googleapis.com/books/v1/volumes?q={book}&key={BooksAPI}&maxResults=1')
    data=response.json()
    if 'items' in data:
        preview=data['items'][0]['volumeInfo'].get('previewLink','Preview not available.')
    else:
        preview='Preview not available.'
    await update.message.reply_text(f'Here is the preview link for the book: {preview}')

async def message_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if context.user_data.get('genre'):
        await book_handler(update,context)
        context.user_data['genre']=False 
    elif context.user_data.get('book'):
        await preview_handler(update,context)
        context.user_data['book']=False 
    elif context.user_data.get('list'):
        context.user_data['list_book']=update.message.text 
        context.user_data['list']=False
        await update.message.reply_text(f'Book "{context.user_data["list_book"]}" selected. Use /reading_list to manage your reading list.')
    else:
        await update.message.reply_text("I am not understanding what you typed...")

if __name__=='__main__':
    print('Starting Bot')
    app=Application.builder().token(Token).build()

    app.add_handler(CommandHandler('start',start_command))
    app.add_handler(CommandHandler('help',help_command))
    app.add_handler(CommandHandler('book',book_command))
    app.add_handler(CommandHandler('preview',preview_command))  
    app.add_handler(CommandHandler('list',list_command))  
    app.add_handler(CommandHandler('reading_list',reading_list_command))    

    app.add_handler(CallbackQueryHandler(reading_list_handler))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND,message_handler))

    print('Polling')
    app.run_polling(poll_interval=2)